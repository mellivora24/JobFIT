import os
import re
import PyPDF2
import pytesseract
from PIL import Image
from docx import Document
from typing import BinaryIO, Union, Optional


def clean_text(text: str) -> str:
    """
    Làm sạch văn bản bằng cách loại bỏ các ký tự đặc biệt và khoảng trắng không cần thiết.
    Args:
        text: Chuỗi văn bản đầu vào
    Returns:
        Chuỗi văn bản đã làm sạch
    """
    if not text:
        return ""

    text = re.sub(r'[\n\r\t]+', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\x20-\x7E\x80-\xFF]', '', text)

    return text.strip()


def extract_text_from_image(image_file: Union[str, BinaryIO]) -> str:
    """
    Trích xuất văn bản từ tệp hình ảnh với các phần tiêu đề và bảng.
    Args:
        image_file: Đường dẫn đến tệp hình ảnh hoặc đối tượng tệp nhị phân
    Returns:
        Văn bản đã trích xuất dưới dạng chuỗi không có ký tự đặc biệt
    """
    # Configure Tesseract with better OCR settings
    custom_config = r'--oem 3 --psm 6'  # Use LSTM OCR Engine + assume single uniform block of text

    # Open image and enhance it for better OCR results
    image = Image.open(image_file)
    # Convert to grayscale for better OCR accuracy
    if image.mode != 'L':
        image = image.convert('L')

    raw_text = pytesseract.image_to_string(image, config=custom_config)
    return clean_text(raw_text)


def extract_text_from_pdf(pdf_file: Union[str, BinaryIO]) -> str:
    """
    Trích xuất văn bản từ tệp PDF với các phần tiêu đề và bảng.
    Args:
        pdf_file: Đường dẫn đến tệp PDF hoặc đối tượng tệp nhị phân
    Returns:
        Văn bản đã trích xuất dưới dạng chuỗi không có ký tự đặc biệt
    """

    reader = PyPDF2.PdfReader(pdf_file)
    text_parts = []

    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
            text_parts.append(clean_text(page_text))
        except Exception as e:
            text_parts.append(f"Error extracting page: {str(e)}")

    return " ".join(filter(None, text_parts))


def extract_text_from_docx(docx_file: Union[str, BinaryIO]) -> str:
    """
    Trích xuất văn bản từ tệp DOCX với các phần tiêu đề và bảng.
    Args:
        docx_file: Đường dẫn đến tệp DOCX hoặc đối tượng tệp nhị phân
    Returns:
        Văn bản đã trích xuất dưới dạng chuỗi không có ký tự đặc biệt
    """

    doc = Document(docx_file)
    text_parts = []

    # Lấy văn bản từ tiêu đề
    for section in doc.sections:
        header = section.header
        if header:
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(clean_text(paragraph.text))
    # Lấy văn bản từ tiêu đề
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(clean_text(paragraph.text))

    # Lấy văn bản từ bảng
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(clean_text(cell.text))
            if row_text:
                text_parts.append(" ".join(row_text))

    return " ".join(filter(None, text_parts))


def extract_text_from_file(file) -> str:
    """
    Trích xuất văn bản từ tệp dựa trên định dạng tệp.
    Hỗ trợ các định dạng: PNG, JPG, JPEG, PDF, DOCX.
    Args:
        file: Tệp đầu vào (có thể là tệp nhị phân hoặc đường dẫn tệp)
    Returns:
        Văn bản đã trích xuất dưới dạng chuỗi không có ký tự đặc biệt
    """
    filename = file.filename.lower() if hasattr(file, 'filename') else str(file).lower()

    extractors = {
        '.png': extract_text_from_image,
        '.jpg': extract_text_from_image,
        '.jpeg': extract_text_from_image,
        ".pdf": extract_text_from_pdf,
        '.docx': extract_text_from_docx,
    }

    _, ext = os.path.splitext(filename)
    ext = re.sub(r"['\">]+", "", ext)

    if ext in extractors:
        return extractors[ext](file)

    supported_formats = ', '.join(extractors.keys())
    raise ValueError(f"File không đúng định dạng ({supported_formats})")


def process_large_file(file_path: str, chunk_size: Optional[int] = None) -> str:
    """
    Process large files efficiently by determining the appropriate extraction method
    and using memory-efficient approaches when needed.

    Args:
        file_path: Path to the file
        chunk_size: Optional size for chunked processing of large files

    Returns:
        Extracted clean text without special characters
    """
    _, ext = os.path.splitext(file_path.lower())

    if ext == '.pdf' and (chunk_size or os.path.getsize(file_path) > 100 * 1024 * 1024):  # > 100MB
        reader = PyPDF2.PdfReader(file_path)
        text_parts = []

        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
                text_parts.append(clean_text(page_text))
            except Exception as e:
                text_parts.append(f"Error extracting page {i}: {str(e)}")

        return " ".join(filter(None, text_parts))

    with open(file_path, 'rb') as f:
        return extract_text_from_file(f)
